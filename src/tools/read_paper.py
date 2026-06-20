"""Tool 2: Read PDF, Word (.docx), Markdown, and plain text papers."""

import os
import re
from src.agent.tool_registry import BaseTool


PAPER_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "data", "papers")


def _get_fitz():
    try:
        import fitz
        return fitz
    except ImportError:
        return None


def _get_pypdf():
    try:
        import pypdf
        return pypdf
    except ImportError:
        return None


def _get_docx():
    try:
        import docx
        return docx
    except ImportError:
        return None


def _detect_format(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    return {
        "": "txt",
        ".txt": "txt",
        ".md": "md",
        ".markdown": "md",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
    }.get(ext, "txt")


def _parse_pdf(path: str) -> str:
    fitz = _get_fitz()
    if fitz:
        doc = fitz.open(path)
        try:
            return "\n".join(page.get_text() for page in doc)
        finally:
            doc.close()
    pypdf = _get_pypdf()
    if pypdf:
        reader = pypdf.PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    return "[PDF library not installed. Install pymupdf or pypdf.]"


def _parse_docx(path: str) -> str:
    docx = _get_docx()
    if not docx:
        return "[python-docx not installed. Install: pip install python-docx]"
    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _parse_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


PARSERS = {"pdf": _parse_pdf, "docx": _parse_docx, "md": _parse_text, "txt": _parse_text}


SECTION_PATTERNS = {
    # Relaxed patterns — no ^ anchor, tolerate leading whitespace/content
    "abstract":    r"\b(abstract|摘要|a b s t r a c t)\b",
    "introduction": r"\b(\d*[\.\s]*(introduction|引言|背景|intro)\b)",
    "related_work": r"\b(\d*[\.\s]*(related work|related works|相关工作|literature review|文献综述)\b)",
    "method":       r"\b(\d*[\.\s]*(method|methods|methodology|approach|framework|model|proposed|方法|模型|框架|our approach|proposed method|算法)\b)",
    "results":      r"\b(\d*[\.\s]*(experiment|experiments|result|results|evaluation|实验|结果|评估|performance|性能)\b)",
    "conclusion":   r"\b(\d*[\.\s]*(conclusion|discussion|limitations|总结|结论|讨论|展望|future work|未来工作)\b)",
    "references":   r"\b(references|bibliography|参考文献|引用文献)\b",
}

# ── header / author / institution noise patterns ──────────────────────
_CONFERENCE_HEADER = re.compile(
    r"(published as a conference paper at|proceedings of the|"
    r"accepted to|accepted at|to appear in|https?://openreview)",
    re.I
)
_INSTITUTION_LINE = re.compile(
    r"(university|institute|college|school of|department of|"
    r"inc\.|ltd\.|laboratory|laboratories|research center|"
    r"universit|école|technische|polytechnique|academy of)",
    re.I
)
_AUTHOR_EMAIL_URL = re.compile(
    r"(arxiv:|doi:|http|@\w+\.\w+|\{.*@.*\}|correspondence|email:|e-mail:)",
    re.I
)
_FOOTER_LINE = re.compile(
    r"^\d+\s*$|^\d+/\d+$|^\d+ of \d+$|^page \d+$|^preprint|"
    r"under review|submitted to|manuscript",
    re.I
)
_COPYRIGHT_LINE = re.compile(
    r"(©|copyright|all rights reserved|licenses/by|cc-by|acm|ieee|springer|elsevier)",
    re.I
)


def _is_header_noise(line: str) -> bool:
    """Check if a line is metadata cruft that should not appear in structured output."""
    s = line.strip()
    if not s or len(s) < 2:
        return True
    if _CONFERENCE_HEADER.search(s):
        return True
    if _AUTHOR_EMAIL_URL.search(s):
        return True
    if _FOOTER_LINE.search(s):
        return True
    if _COPYRIGHT_LINE.search(s):
        return True
    # Short line that's just institution numbers like "1,2,3" or "1 2 3"
    if re.match(r"^[\d,\s\*†‡§¶∥]+$", s) and len(s) < 30:
        return True
    return False


def _is_likely_author_line(line: str) -> bool:
    """Heuristic: line looks like author names + affiliations."""
    s = line.strip()
    # Contains superscript numbers/asterisks typical of author blocks
    if re.search(r"[\d\*†‡]+\s*[,\s]|,\s*[\d\*†‡]+", s):
        if len(s) > 20 and len(s) < 300:
            return True
    # Institution line with department markers
    if _INSTITUTION_LINE.search(s) and len(s) < 200:
        return True
    return False


def _find_section_pos(lower: str, pattern: str) -> int | None:
    """Find a section pattern in a line; returns match start pos or None."""
    m = re.search(pattern, lower, re.I)
    if m and m.start() <= 6:  # section header must be near line start
        return m.start()
    return None


def _extract_sections(text: str) -> dict:
    sections = {
        "title": "",
        "abstract": "",
        "introduction": "",
        "related_work": "",
        "method": "",
        "results": "",
        "conclusion": "",
        "references": "",
    }
    if not text:
        return sections

    lines = [line.strip() for line in text.splitlines()]
    visible = [line for line in lines if line]

    # ── Phase 1: find explicit section boundaries ──────────────────
    section_starts: list[tuple[int, str]] = []  # (line_index, section_key)
    for idx, line in enumerate(visible):
        lower = line.lower().strip()
        for key, pattern in SECTION_PATTERNS.items():
            pos = _find_section_pos(lower, pattern)
            if pos is not None:
                section_starts.append((idx, key))
                break  # first matching section wins for this line

    # ── Phase 2: extract title (up to first author/institution-free block) ──
    title_end = section_starts[0][0] if section_starts else min(8, len(visible))
    # Try to find a better title end: skip author/inst lines, stop at abstract-like text
    title_lines = []
    in_author_block = False
    for idx in range(min(15, len(visible))):
        line = visible[idx]
        if _is_header_noise(line):
            continue
        if _is_likely_author_line(line):
            in_author_block = True
            continue
        if in_author_block:
            # After author block, check if this is an abstract section start or real text
            lower = line.lower().strip()
            if _find_section_pos(lower, SECTION_PATTERNS["abstract"]) is not None:
                break
            # A long line (> 100 chars) after author block → likely abstract
            if len(line) > 100:
                title_lines.append(line[:200])
                break
            # Short non-author line after author block → back to title/header
            if len(line) < 80:
                in_author_block = False
        if not in_author_block and not _is_likely_author_line(line):
            title_lines.append(line)
        if len("\n".join(title_lines)) > 800:
            break

    sections["title"] = "\n".join(title_lines[:8])[:2000]

    # ── Phase 3: walk lines with detected sections ──────────────────
    # Build a map of line_idx → section_key
    section_map = {}
    for idx, key in section_starts:
        section_map[idx] = key

    # If no explicit abstract section found, find the first long paragraph after title
    abstract_start = None
    if not any(k == "abstract" for _, k in section_starts):
        for idx in range(len(title_lines) + 1, min(len(visible), 30)):
            line = visible[idx]
            if _is_header_noise(line) or _is_likely_author_line(line):
                continue
            if len(line) > 120:  # first substantial paragraph
                abstract_start = idx
                break

    current = "abstract"
    for idx, line in enumerate(visible):
        # Skip lines already consumed by title
        if idx < len(title_lines):
            continue
        # If abstract not found yet and we're before abstract_start, skip
        if abstract_start is not None and current == "abstract" and idx < abstract_start:
            continue

        lower = line.lower().strip()

        # Check section boundary
        if idx in section_map:
            current = section_map[idx]
            continue

        # Skip noise
        if _is_header_noise(line) and len(line) < 80:
            continue
        if _is_likely_author_line(line) and current in ("abstract", "title"):
            continue

        if current in sections:
            prev = sections.get(current, "")
            sections[current] = ((prev + "\n") if prev else "") + line
            sections[current] = sections[current][:3000]

    # ── Phase 4: clean up — if abstract is empty, grab first 2000 chars from raw ──
    if not sections["abstract"].strip():
        # Use raw text after the title area
        raw_start = min(10, len(visible))
        fallback = "\n".join(visible[raw_start:])[:2000]
        sections["abstract"] = fallback

    return sections


def translate_text(text: str, target_lang: str = "zh") -> str:
    return f"[Translation to {target_lang} requested. This is handled by the chat interface with the active LLM model.]"


class ReadPaperTool(BaseTool):
    name = "read_paper"
    description = (
        "Read and extract structured information from a local PDF, Word (.docx), "
        "Markdown, or text file. Returns title, abstract, method, results, "
        "conclusion, and references by section."
    )
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Path to the document file on disk"},
            "output": {
                "type": "string",
                "enum": ["full", "sections"],
                "description": "full=entire text, sections=structured by section",
                "default": "sections",
            },
        },
        "required": ["path"],
    }

    def run(self, path: str = "", output: str = "sections") -> str:
        if not os.path.exists(path):
            alt = os.path.join(PAPER_DIR, os.path.basename(path))
            if os.path.exists(alt):
                path = alt
            else:
                return f"File not found: {path}"

        fmt = _detect_format(path)
        parser = PARSERS.get(fmt)
        if not parser:
            return f"Unsupported format: {fmt}"

        raw = parser(path)
        if raw.startswith("["):
            return raw

        lines = [f"File: {os.path.basename(path)}", f"Format: {fmt.upper()}", f"Size: {len(raw)} chars", ""]
        if output == "full":
            lines.append(raw[:8000])
            return "\n".join(lines)

        sections = _extract_sections(raw)
        labels = {
            "title": "Title",
            "abstract": "Abstract",
            "introduction": "Introduction",
            "related_work": "Related Work",
            "method": "Method / Approach",
            "results": "Experiments / Results",
            "conclusion": "Conclusion / Discussion",
            "references": "References",
        }
        for key, label in labels.items():
            value = sections.get(key, "").strip()
            if value:
                lines.append(f"[{label}]")
                lines.append(value[:1800])
                lines.append("")

        if len(lines) <= 4:
            lines.append("[No clear sections detected - showing raw content]")
            lines.append(raw[:3000])
        return "\n".join(lines)
