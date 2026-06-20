# -*- coding: utf-8 -*-
"""Literature research workspace."""

import hashlib
import html
import json
import os
import pathlib
import re
import sys
from datetime import datetime

import streamlit as st

sys.path.insert(0, str(pathlib.Path(__file__).parent.parent))

from src.agent.core import AcademicLitAgent
from src.agent.llm_client import LLMClient
from src.agent.tool_registry import ToolRegistry
from src.tools import (
    CitationTraverserTool,
    CodeExecutorTool,
    LiteratureComparatorTool,
    ReadPaperTool,
    ResearchGapAnalyzerTool,
    SQLitePaperDBTool,
    SearchPapersTool,
)
from src.tools.read_paper import PARSERS, _detect_format, _extract_sections


PAPER_DIR = pathlib.Path("data/papers")
TOOL_REGISTRY_VERSION = "2026-06-17.4"
NO_LLM = "无需 LLM"
AUTO_MODEL = "自动选择"


def ensure_tool_registry() -> None:
    if (
        "tool_registry" in st.session_state
        and st.session_state.get("tool_registry_version") == TOOL_REGISTRY_VERSION
    ):
        return
    registry = ToolRegistry()
    for tool in (
        SearchPapersTool(),
        ReadPaperTool(),
        SQLitePaperDBTool(),
        LiteratureComparatorTool(),
        ResearchGapAnalyzerTool(),
        CitationTraverserTool(),
        CodeExecutorTool(),
    ):
        registry.register(tool)
    st.session_state.tool_registry = registry
    st.session_state.tool_registry_version = TOOL_REGISTRY_VERSION


def active_model_names() -> list[str]:
    return [m.name for m in st.session_state.model_manager.list_models()]


def model_options(include_auto: bool = False, include_no_llm: bool = False) -> list[str]:
    options = active_model_names()
    if include_auto:
        options = [AUTO_MODEL] + options
    if include_no_llm:
        options = options + [NO_LLM]
    return options or ([NO_LLM] if include_no_llm else [])


def normalize_model_choice(choice: str | None) -> str:
    models = active_model_names()
    if choice in models:
        return choice
    if st.session_state.get("current_model") in models:
        return st.session_state.current_model
    return models[0] if models else ""


def model_index(options: list[str], value: str | None) -> int:
    if value in options:
        return options.index(value)
    return 0


def get_llm_client(model_name: str | None = None) -> LLMClient | None:
    name = normalize_model_choice(model_name)
    config = st.session_state.model_manager.get_model(name)
    if not config:
        return None
    return LLMClient(
        api_key=config.api_key or "",
        base_url=config.base_url or "",
        model_name=config.model_name or "",
    )


def get_agent(model_name: str | None = None) -> AcademicLitAgent | None:
    client = get_llm_client(model_name)
    if not client:
        return None
    return AcademicLitAgent(llm_client=client, tool_registry=st.session_state.tool_registry)


def page_css() -> None:
    st.markdown(
        """
        <style>
        /* 彻底隐藏 Streamlit 顶部系统栏：Deploy、三点菜单、顶部装饰线 */
        div[data-testid="stToolbar"],
        div[data-testid="stDecoration"],
        div[data-testid="stStatusWidget"],
        #MainMenu,
        footer {
            display: none !important;
            visibility: hidden !important;
            height: 0 !important;
            min-height: 0 !important;
            max-height: 0 !important;
        }
        header[data-testid="stHeader"] {
            height: 0 !important;
            min-height: 0 !important;
            background: transparent !important;
            pointer-events: none;
        }
        header[data-testid="stHeader"] * {
            pointer-events: auto;
        }
        [data-testid="stSidebarCollapseButton"],
        [data-testid="stSidebarCollapsedControl"],
        button[title="Open sidebar"],
        button[title="Close sidebar"] {
            display: flex !important;
            visibility: visible !important;
            opacity: 1 !important;
            z-index: 100000 !important;
        }
        [data-testid="stSidebarCollapseButton"] {
            position: fixed !important;
            left: 302px !important;
            top: 10px !important;
            width: 34px !important;
            height: 34px !important;
            align-items: center !important;
            justify-content: center !important;
            border-radius: 10px !important;
            background: #ffffff !important;
            border: 1px solid #dde5f0 !important;
            box-shadow: 0 8px 20px rgba(15, 23, 42, 0.08) !important;
        }
        [data-testid="stSidebarCollapseButton"] button {
            width: 32px !important;
            height: 32px !important;
        }
        [data-testid="stSidebarCollapsedControl"] {
            position: fixed !important;
            left: 18px !important;
            top: 18px !important;
            width: 42px !important;
            height: 42px !important;
            border-radius: 12px !important;
            background: #ffffff !important;
            border: 1px solid #dce3ee !important;
            box-shadow: 0 12px 30px rgba(15, 23, 42, 0.14) !important;
        }

        /* 去掉 Streamlit 给顶部预留的空白 */
        [data-testid="stAppViewContainer"],
        [data-testid="stAppViewBlockContainer"],
        section.main > div,
        .block-container {
            padding-top: 0 !important;
            margin-top: 0 !important;
        }

        /* 页面主体继续保持宽屏铺满 */
        .block-container {
            padding-top: 0 !important;
            padding-bottom: 112px;
            max-width: 1480px;
        }
        [data-testid="stSidebar"] {
            background: #ffffff;
            border-right: 1px solid #dfe5ee;
        }
        [data-testid="stSidebar"] > div:first-child {
            padding-top: 28px;
        }
        [data-testid="stSidebarUserContent"] {
            margin-top: -110px;
        }
        [data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
            gap: 0.75rem;
        }
        [data-testid="stSidebarNav"] {
            display: none !important;
        }
        .sidebar-brand {
            display: flex;
            align-items: center;
            gap: 14px;
            padding: 2px 0 18px;
            margin-bottom: 18px;
            border-bottom: 1px solid #e4e9f1;
        }
        .sidebar-brand-mark {
            width: 42px;
            height: 42px;
            border-radius: 12px;
            background: #3d6ee8;
            color: #ffffff;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 25px;
            font-weight: 900;
            line-height: 1;
            box-shadow: 0 8px 18px rgba(61, 110, 232, 0.22);
        }
        .sidebar-brand-title {
            margin: 0;
            color: #0f172a;
            font-size: 18px;
            font-weight: 850;
            line-height: 1.2;
        }
        .sidebar-brand-subtitle {
            margin: 5px 0 0;
            color: #667085;
            font-size: 13px;
            line-height: 1.2;
        }
        h1, h2, h3 { letter-spacing: 0; }
        div[data-testid="stMetric"] {
            background: #ffffff;
            border: 1px solid #ebeff5;
            border-radius: 8px;
            padding: 10px 12px;
            box-shadow: 0 4px 18px rgba(30, 41, 59, 0.03);
        }
        div[data-testid="stMetric"] label {
            color: #697386;
            font-size: 12px;
        }
        div[data-testid="stMetricValue"] {
            font-size: 21px;
            font-weight: 780;
        }
        .topbar-title h1 {
            margin: 0;
            font-size: 30px;
            line-height: 1.2;
            color: #0f172a;
            font-weight: 900;
        }
        .topbar-title p {
            margin: 5px 0 0;
            color: #697386;
            font-size: 14px;
        }
        .topbar-control-label {
            display: flex;
            align-items: center;
            gap: 8px;
            height: 38px;
            padding: 0 13px;
            margin-top: 1px;
            border: 1px solid #dce3ee;
            border-radius: 8px;
            background: #ffffff;
            color: #344054;
            font-size: 13px;
            font-weight: 750;
            white-space: nowrap;
            box-shadow: 0 3px 12px rgba(15, 23, 42, 0.04);
        }
        .topbar-control-label .status-dot {
            flex: 0 0 auto;
        }
        .topbar-select [data-testid="stSelectbox"] label {
            display: none !important;
        }
        .topbar-select div[data-baseweb="select"] {
            min-height: 38px !important;
            height: 38px !important;
            border-radius: 8px !important;
            border: 1px solid #dce3ee !important;
            background: #ffffff !important;
            box-shadow: 0 3px 12px rgba(15, 23, 42, 0.04) !important;
        }
        .topbar-select div[data-baseweb="select"] > div {
            padding-left: 0 !important;
        }
        button[data-testid="stBaseButton-primary"] {
            min-height: 38px !important;
            height: 38px !important;
            border-radius: 8px !important;
            background: #2f6feb !important;
            border: 1px solid #245bd4 !important;
            color: #ffffff !important;
            font-weight: 800 !important;
            white-space: nowrap !important;
            box-shadow: 0 2px 0 rgba(20, 58, 138, 0.35) !important;
        }
        button[data-testid="stBaseButton-secondary"] {
            min-height: 38px !important;
            height: 38px !important;
            border-radius: 8px !important;
            font-weight: 760 !important;
            white-space: nowrap !important;
        }
        .status-dot {
            width: 8px;
            height: 8px;
            border-radius: 50%;
            background: #17b26a;
            display: inline-block;
        }
        .primary-action {
            height: 38px;
            border-radius: 7px;
            background: #1f6feb;
            color: #ffffff;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 13px;
            font-weight: 700;
            white-space: nowrap;
            margin-top: 28px;
        }
        .section-card {
            background: #ffffff;
            border: 1px solid #ebeff5;
            border-radius: 8px;
            box-shadow: 0 16px 40px rgba(30, 41, 59, 0.08);
            overflow: hidden;
        }
        div[data-testid="stForm"] {
            position: fixed;
            left: max(330px, calc((100vw - 1480px) / 2 + 300px + 100px));
            bottom: 16px;
            width: min(512px, calc((100vw - 360px) * 0.475));
            height: auto !important;
            min-height: 70px !important;
            max-height: 190px !important;
            overflow: visible !important;
            z-index: 999;
            border: 1px solid #dfe3ea !important;
            border-radius: 26px !important;
            padding: 10px 14px !important;
            background: #ffffff !important;
            box-shadow: 0 18px 44px rgba(15, 23, 42, 0.12) !important;
        }
        div[data-testid="stForm"] [data-testid="column"] {
            display: flex;
            align-items: center;
            justify-content: center;
        }
        div[data-testid="stForm"] [data-testid="stVerticalBlock"],
        div[data-testid="stForm"] [data-testid="stHorizontalBlock"] {
            gap: 0.35rem !important;
            min-height: 0 !important;
            align-items: center !important;
        }
        .composer-plus {
            height: 46px;
            width: 28px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 27px;
            line-height: 1;
            color: #667085;
        }
        div[data-testid="stForm"] div[data-testid="stTextArea"] {
            width: 100%;
        }
        div[data-testid="stForm"] div[data-baseweb="textarea"],
        div[data-testid="stForm"] div[data-baseweb="textarea"]:focus-within {
            border: none !important;
            box-shadow: none !important;
            outline: none !important;
            background: transparent !important;
        }
        div[data-testid="stForm"] div[data-testid="stTextArea"] textarea {
            height: auto !important;
            min-height: 48px !important;
            max-height: 132px !important;
            field-sizing: content;
            resize: none !important;
            overflow-y: auto !important;
            border: none !important;
            box-shadow: none !important;
            background: #f1f3f6 !important;
            border-radius: 13px !important;
            padding: 13px 15px !important;
            font-size: 14px !important;
            line-height: 1.45 !important;
        }
        div[data-testid="stForm"] div[data-testid="stTextArea"] textarea:focus {
            border: none !important;
            box-shadow: none !important;
            outline: none !important;
        }
        div[data-testid="stForm"] [data-testid="InputInstructions"] {
            display: none !important;
        }
        div[data-testid="stForm"] button[kind="primaryFormSubmit"] {
            border-radius: 12px !important;
            min-width: 70px !important;
            width: 70px !important;
            height: 48px !important;
            padding: 0 14px !important;
            background: #2f6feb !important;
            border: 1px solid #245bd4 !important;
            color: #ffffff !important;
            font-weight: 800 !important;
            box-shadow: 0 2px 0 rgba(20, 58, 138, 0.45) !important;
        }
        div[data-testid="stForm"] button[kind="secondaryFormSubmit"] {
            border-radius: 50% !important;
            min-width: 38px !important;
            width: 38px !important;
            height: 38px !important;
            padding: 0 !important;
            font-size: 22px !important;
        }
        div[data-testid="stForm"] div[data-baseweb="select"] {
            min-height: 48px !important;
            height: 48px !important;
            border-radius: 13px !important;
            background: #f1f3f6 !important;
        }
        div[data-testid="stForm"] div[data-baseweb="select"] span {
            font-size: 13px !important;
        }
        .scroll-paper {
            height: 640px;
            overflow-y: auto;
            background: #f0f2f6;
            border-radius: 8px;
            padding: 14px;
            border: 1px solid #e5eaf2;
        }
        .translation-scroll {
            height: 640px;
            overflow-y: auto;
            border: 1px solid #e5eaf2;
            border-radius: 8px;
            padding: 14px 16px;
            background: #ffffff;
            line-height: 1.7;
        }
        .panel-head {
            height: 48px;
            border-bottom: 1px solid #ebeff5;
            padding: 0 14px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 12px;
            background: #ffffff;
        }
        .panel-head h3 {
            margin: 0;
            font-size: 15px;
            color: #192232;
        }
        .panel-head span {
            color: #697386;
            font-size: 12px;
        }
        .tool-card {
            border: 1px solid #dfe5ee;
            border-radius: 8px;
            padding: 10px;
            background: #ffffff;
            min-height: 92px;
        }
        .tool-card strong {
            display: block;
            font-size: 13px;
            color: #192232;
            margin-bottom: 4px;
        }
        .tool-card p {
            min-height: 32px;
            margin: 0 0 8px;
            color: #697386;
            font-size: 12px;
            line-height: 1.35;
        }
        .message-card {
            border: 1px solid #ebeff5;
            border-radius: 8px;
            padding: 12px;
            margin-bottom: 12px;
            background: #ffffff;
        }
        .message-card.user {
            background: #eef4ff;
            border-color: #c9ddff;
            margin-left: 9%;
        }
        .message-card.assistant {
            margin-right: 9%;
        }
        .message-meta {
            display: flex;
            align-items: center;
            gap: 8px;
            color: #8a95a6;
            font-size: 11px;
            margin-bottom: 7px;
            flex-wrap: wrap;
        }
        .model-badge {
            border: 1px solid #cdd8ea;
            background: #f8fbff;
            border-radius: 999px;
            padding: 2px 7px;
            color: #2459a6;
            font-weight: 700;
            font-size: 11px;
        }
        .paper-chip {
            border: 1px solid #ebeff5;
            background: #f9fafc;
            border-radius: 8px;
            padding: 9px 10px;
            margin-bottom: 8px;
        }
        .paper-chip.active {
            border-color: #9bbcf5;
            background: #f5f9ff;
        }
        .paper-title {
            color: #192232;
            font-size: 13px;
            font-weight: 760;
            line-height: 1.3;
        }
        .paper-meta {
            margin-top: 5px;
            display: flex;
            justify-content: space-between;
            color: #8a95a6;
            font-size: 11px;
        }
        .tag {
            border-radius: 999px;
            padding: 3px 7px;
            background: #eef4ff;
            color: #2459a6;
            font-size: 11px;
            font-weight: 650;
            display: inline-block;
            margin: 6px 4px 0 0;
        }
        .tag.teal {
            background: #e9f7f4;
            color: #087968;
        }
        .tag.amber {
            background: #fff3dc;
            color: #915c00;
        }
        .reader-shell {
            background: #f0f2f6;
            border-radius: 8px;
            padding: 14px;
            min-height: 560px;
        }
        .inspector-card {
            border: 1px solid #ebeff5;
            border-radius: 8px;
            padding: 10px;
            margin-bottom: 10px;
            background: #fbfcfe;
        }
        .inspector-card h4 {
            margin: 0 0 8px;
            font-size: 13px;
            color: #192232;
        }
        .kv-row {
            display: grid;
            grid-template-columns: 56px 1fr;
            gap: 8px;
            font-size: 12px;
            color: #697386;
            margin-bottom: 6px;
            line-height: 1.35;
        }
        .kv-row b {
            color: #344054;
        }
        .hint-box {
            border-left: 3px solid #af6b00;
            background: #fff8ea;
            padding: 7px 8px;
            border-radius: 5px;
            color: #6f4700;
            font-size: 12px;
            margin-bottom: 7px;
            line-height: 1.4;
        }
        .hint-box.teal {
            border-left-color: #118b7a;
            background: #eefaf7;
            color: #0f665b;
        }
        .stButton > button, .stDownloadButton > button {
            border-radius: 7px;
        }
        div[data-testid="stTextInput"] input,
        div[data-testid="stSelectbox"] div[data-baseweb="select"] {
            border-radius: 7px;
        }
        div[data-testid="stSelectbox"] div[data-baseweb="select"] span {
            font-size: 13px !important;
        }
        /* workspace_v3 compact visual pass */
        .stApp {
            background: #f5f7fb !important;
            color: #192232 !important;
            font-family: "Inter", "Segoe UI", "Microsoft YaHei", system-ui, sans-serif !important;
        }
        .block-container {
            max-width: 1480px !important;
            padding-left: 22px !important;
            padding-right: 22px !important;
            padding-bottom: 96px !important;
        }
        [data-testid="stSidebar"] {
            background: #ffffff !important;
            border-right: 1px solid #dfe5ee !important;
        }
        [data-testid="stSidebarCollapseButton"] {
            left: 248px !important;
            top: 18px !important;
            width: 30px !important;
            height: 30px !important;
            border-radius: 8px !important;
            box-shadow: none !important;
        }
        [data-testid="stSidebarCollapseButton"] button {
            width: 28px !important;
            height: 28px !important;
        }
        body:has([data-testid="stSidebar"].st-emotion-cache-1udkqym) [data-testid="stSidebarCollapseButton"] {
            left: 310px !important;
            top: 10px !important;
            width: 34px !important;
            height: 34px !important;
            box-shadow: 0 8px 20px rgba(15, 23, 42, 0.08) !important;
        }
        [data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
            gap: 0.68rem !important;
        }
        [data-testid="stSidebar"] h3 {
            margin: 0 0 8px !important;
            color: #475467 !important;
            font-size: 13px !important;
            line-height: 1.25 !important;
            font-weight: 780 !important;
        }
        [data-testid="stSidebar"] hr {
            margin: 10px 0 !important;
            border-color: #ebeff5 !important;
        }
        [data-testid="stSidebar"] small,
        [data-testid="stSidebar"] p,
        [data-testid="stSidebar"] label {
            font-size: 12px !important;
            color: #697386 !important;
        }
        [data-testid="stSidebar"] [data-testid="stFileUploader"] section {
            border: 1px dashed #adc2df !important;
            border-radius: 8px !important;
            background: #f7fbff !important;
            padding: 10px 12px !important;
        }
        [data-testid="stSidebar"] [data-testid="stFileUploader"] button {
            height: 30px !important;
            min-height: 30px !important;
            border-radius: 6px !important;
            font-size: 12px !important;
        }
        [data-testid="stSidebar"] [data-testid="stFileUploader"] [data-testid="stMarkdownContainer"] {
            color: #2459a6 !important;
            font-size: 13px !important;
            font-weight: 650 !important;
        }
        .asset-title-row {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 10px;
            margin: 2px 0 10px;
        }
        .asset-title-row h3 {
            margin: 0 !important;
            color: #344054 !important;
            font-size: 16px !important;
            font-weight: 850 !important;
            line-height: 1.2 !important;
        }
        .asset-add {
            width: 32px;
            height: 32px;
            border: 1px solid #dfe5ee;
            border-radius: 8px;
            display: flex;
            align-items: center;
            justify-content: center;
            color: #344054;
            background: #ffffff;
            font-size: 20px;
            line-height: 1;
        }
        .library-card {
            border: 1px solid #dfe5ee;
            border-radius: 10px;
            background: #ffffff;
            padding: 12px 12px 10px;
            margin-bottom: 10px;
        }
        .library-card.active {
            border-color: #9bbcf5;
            background: #f5f9ff;
        }
        .library-card-head {
            display: grid;
            grid-template-columns: minmax(0, 1fr) auto;
            gap: 8px;
            align-items: start;
        }
        .library-card-title {
            color: #0f172a;
            font-size: 15px;
            font-weight: 850;
            line-height: 1.35;
            word-break: break-word;
        }
        .library-menu-placeholder {
            width: 34px;
            height: 30px;
        }
        .library-meta-row {
            margin-top: 8px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 8px;
            color: #8a95a6;
            font-size: 13px;
        }
        .library-tags {
            margin-top: 10px;
            display: flex;
            gap: 8px;
            flex-wrap: wrap;
        }
        .library-menu-wrap {
            margin-top: -104px;
            display: flex;
            justify-content: flex-end;
            pointer-events: none;
        }
        .library-menu-wrap [data-testid="stPopover"] {
            pointer-events: auto;
        }
        [data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] {
            border-color: #dfe5ee !important;
            border-radius: 10px !important;
            background: #ffffff !important;
        }
        [data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"]:has(.library-card-inner.active) {
            border-color: #9bbcf5 !important;
            background: #f5f9ff !important;
        }
        [data-testid="stSidebar"] [data-testid="stPopover"] button {
            min-height: 30px !important;
            height: 30px !important;
            border-radius: 6px !important;
            padding: 0 8px !important;
            font-size: 16px !important;
            line-height: 1 !important;
        }
        .library-card-inner {
            padding: 0;
        }
        .library-card-inner .library-card-title {
            margin-top: 2px;
        }
        .library-menu-wrap button {
            width: 34px !important;
            min-width: 34px !important;
            height: 30px !important;
            min-height: 30px !important;
            padding: 0 !important;
            border-radius: 6px !important;
            font-size: 18px !important;
            font-weight: 800 !important;
        }
        .sidebar-brand {
            gap: 10px !important;
            padding: 2px 4px 10px !important;
            margin-bottom: 14px !important;
            border-bottom: 1px solid #ebeff5 !important;
        }
        .sidebar-brand-mark {
            width: 34px !important;
            height: 34px !important;
            border-radius: 8px !important;
            background: #1f6feb !important;
            font-size: 17px !important;
            font-weight: 800 !important;
            box-shadow: none !important;
        }
        .sidebar-brand-title {
            font-size: 15px !important;
            line-height: 1.25 !important;
            font-weight: 800 !important;
        }
        .sidebar-brand-subtitle {
            margin-top: 2px !important;
            font-size: 12px !important;
            color: #697386 !important;
        }
        .topbar-title h1 {
            font-size: 22px !important;
            line-height: 1.25 !important;
            font-weight: 850 !important;
            color: #192232 !important;
        }
        .topbar-title {
            padding-left: 22px !important;
        }
        .topbar-title p {
            margin-top: 4px !important;
            font-size: 13px !important;
            color: #697386 !important;
            line-height: 1.35 !important;
        }
        div[data-testid="stMetric"] {
            border: 1px solid #ebeff5 !important;
            border-radius: 8px !important;
            padding: 12px !important;
            box-shadow: 0 4px 18px rgba(30, 41, 59, 0.03) !important;
            min-height: 108px !important;
        }
        div[data-testid="stMetric"] label,
        div[data-testid="stMetric"] [data-testid="stMetricDelta"] {
            font-size: 12px !important;
            color: #697386 !important;
        }
        div[data-testid="stMetricValue"] {
            font-size: 20px !important;
            line-height: 1.15 !important;
            font-weight: 780 !important;
        }
        .section-card {
            border: 1px solid #ebeff5 !important;
            border-radius: 8px !important;
            box-shadow: 0 16px 40px rgba(30, 41, 59, 0.08) !important;
            background: #ffffff !important;
        }
        .panel-head {
            height: 48px !important;
            padding: 0 14px !important;
            border-bottom: 1px solid #ebeff5 !important;
            background: #ffffff !important;
        }
        .panel-head h3 {
            font-size: 15px !important;
            font-weight: 760 !important;
        }
        .panel-head span {
            font-size: 12px !important;
            color: #697386 !important;
        }
        .tool-card {
            min-height: 84px !important;
            padding: 10px !important;
            border-radius: 8px !important;
            border-color: #dfe5ee !important;
            box-shadow: none !important;
        }
        .tool-card strong,
        .message-card h1,
        .message-card h2,
        .message-card h3,
        .message-card h4 {
            font-size: 13px !important;
        }
        .tool-card p {
            min-height: 30px !important;
            margin-bottom: 7px !important;
            font-size: 11px !important;
            line-height: 1.35 !important;
        }
        .tool-card div[data-baseweb="select"] {
            min-height: 24px !important;
            height: 24px !important;
            border-radius: 6px !important;
            background: #f8fafc !important;
        }
        .tool-card div[data-baseweb="select"] span {
            font-size: 11px !important;
            font-weight: 650 !important;
        }
        .message-card {
            max-width: 88% !important;
            border-radius: 8px !important;
            padding: 10px 12px !important;
            margin-bottom: 12px !important;
            font-size: 13px !important;
            line-height: 1.55 !important;
            color: #192232 !important;
        }
        .message-card.user {
            margin-left: 12% !important;
            background: #eef4ff !important;
            border-color: #c9ddff !important;
        }
        .message-card.assistant {
            margin-right: 12% !important;
            background: #ffffff !important;
        }
        .message-meta,
        .model-badge {
            font-size: 11px !important;
        }
        .paper-chip {
            padding: 8px 9px !important;
            border-radius: 7px !important;
            margin-bottom: 7px !important;
            background: #ffffff !important;
        }
        .paper-chip.active {
            background: #eef4ff !important;
            border-color: #9bbcf5 !important;
        }
        .paper-title {
            font-size: 12px !important;
            font-weight: 740 !important;
        }
        .paper-meta {
            font-size: 11px !important;
        }
        .tag {
            font-size: 11px !important;
            padding: 3px 7px !important;
            margin-top: 6px !important;
        }
        .reader-shell,
        .scroll-paper,
        .translation-scroll {
            border-radius: 8px !important;
            background: #f0f2f6 !important;
        }
        .inspector-card {
            padding: 10px !important;
            border-radius: 8px !important;
            font-size: 12px !important;
        }
        .inspector-card h4 {
            font-size: 12px !important;
        }
        .kv-row,
        .hint-box {
            font-size: 11px !important;
        }
        .stButton > button,
        .stDownloadButton > button {
            min-height: 30px !important;
            height: 30px !important;
            border-radius: 6px !important;
            padding: 0 9px !important;
            font-size: 12px !important;
            font-weight: 650 !important;
        }
        button[data-testid="stBaseButton-primary"] {
            min-height: 34px !important;
            height: 34px !important;
            border-radius: 7px !important;
            background: #1f6feb !important;
            border-color: #1557c0 !important;
            font-size: 13px !important;
            font-weight: 700 !important;
            box-shadow: none !important;
        }
        div[data-testid="stSelectbox"] div[data-baseweb="select"] {
            min-height: 34px !important;
            height: 34px !important;
            border-radius: 7px !important;
            border-color: #dfe5ee !important;
            background: #ffffff !important;
        }
        div[data-testid="stSelectbox"] div[data-baseweb="select"] span {
            font-size: 12px !important;
        }
        div[data-testid="stForm"] {
            position: fixed !important;
            left: max(322px, calc((100vw - 1480px) / 2 + 322px)) !important;
            right: auto !important;
            bottom: 10px !important;
            width: min(576px, calc((100vw - 360px) * 0.533)) !important;
            min-height: 58px !important;
            max-height: none !important;
            margin: 0 !important;
            padding: 10px 12px !important;
            border: 0 !important;
            border: 1px solid #ebeff5 !important;
            border-radius: 0 !important;
            background: #fbfcfe !important;
            box-shadow: none !important;
            z-index: 999 !important;
        }
        body:has([data-testid="stSidebar"].st-emotion-cache-1udkqym) div[data-testid="stForm"] {
            left: 22px !important;
            width: min(738px, calc(100vw - 680px)) !important;
        }
        div[data-testid="stForm"] [data-testid="stHorizontalBlock"] {
            gap: 0.5rem !important;
        }
        .composer-plus {
            width: 30px !important;
            height: 38px !important;
            margin-top: -8px !important;
            font-size: 24px !important;
            color: #667085 !important;
        }
        div[data-testid="stForm"] div[data-testid="stTextArea"] textarea {
            min-height: 38px !important;
            height: 38px !important;
            max-height: 118px !important;
            border: 1px solid #dfe5ee !important;
            border-radius: 7px !important;
            background: #ffffff !important;
            padding: 9px 12px !important;
            font-size: 13px !important;
            line-height: 1.35 !important;
        }
        div[data-testid="stForm"] div[data-baseweb="select"] {
            min-height: 38px !important;
            height: 38px !important;
            border-radius: 7px !important;
            border: 1px solid #dfe5ee !important;
            background: #ffffff !important;
        }
        div[data-testid="stForm"] div[data-baseweb="select"] span {
            font-size: 12px !important;
            font-weight: 650 !important;
        }
        div[data-testid="stForm"] button[kind="primaryFormSubmit"] {
            width: auto !important;
            min-width: 58px !important;
            height: 38px !important;
            border-radius: 7px !important;
            padding: 0 12px !important;
            background: #1f6feb !important;
            border-color: #1557c0 !important;
            font-size: 13px !important;
            font-weight: 700 !important;
            box-shadow: none !important;
        }
        @media (max-width: 1100px) {
            div[data-testid="stForm"] {
                left: 16px !important;
                right: 16px !important;
                width: auto !important;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def add_uploaded_file(uploaded) -> None:
    PAPER_DIR.mkdir(parents=True, exist_ok=True)
    path = PAPER_DIR / uploaded.name
    path.write_bytes(uploaded.getbuffer())
    source_hash = file_sha256(path)
    page_count = estimate_page_count(path)
    record = {
        "name": uploaded.name,
        "path": str(path),
        "timestamp": datetime.now().strftime("%m-%d %H:%M"),
        "paper_id": f"local:{source_hash[:16]}",
        "page_count": page_count,
        "favorite": False,
        "tags": [],
        "status": "刚刚解析",
    }
    saved = save_paper_record_to_db(record, source_hash)
    if saved.get("paper_id"):
        record["paper_id"] = saved["paper_id"]
    existing = next((p for p in st.session_state.uploaded_papers if p["name"] == uploaded.name), None)
    if existing:
        existing.update(record)
    else:
        st.session_state.uploaded_papers.append(record)
    st.session_state.current_doc = str(path)
    st.session_state.current_doc_name = uploaded.name
    st.session_state.doc_translation = ""


def file_sha256(path: pathlib.Path) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def estimate_page_count(path: pathlib.Path) -> int:
    fmt = _detect_format(str(path))
    try:
        if fmt == "pdf" and _has_lib("fitz"):
            import fitz

            doc = fitz.open(str(path))
            try:
                return len(doc)
            finally:
                doc.close()
        if fmt == "docx" and _has_lib("docx"):
            import docx

            doc = docx.Document(str(path))
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            return max(1, round(len(paragraphs) / 8))
    except Exception:
        return 0
    return 0


def save_paper_record_to_db(record: dict, source_hash: str = "") -> dict:
    out = st.session_state.tool_registry.execute(
        {
            "tool": "sqlite_paper_db",
            "kwargs": {
                "operation": "save",
                "paper_id": record.get("paper_id", ""),
                "title": record.get("name", ""),
                "year": str(record.get("year", "")),
                "file_path": record.get("path", ""),
                "source_hash": source_hash,
                "page_count": str(record.get("page_count", 0)),
                "status": record.get("status", ""),
                "tags": ",".join(record.get("tags", [])),
            },
        }
    )
    try:
        return json.loads(out)
    except json.JSONDecodeError:
        return {}


def paper_library_records() -> list[dict]:
    out = st.session_state.tool_registry.execute({"tool": "sqlite_paper_db", "kwargs": {"operation": "list_json"}})
    try:
        records = json.loads(out)
    except json.JSONDecodeError:
        records = []
    if not records:
        return []
    return [normalize_library_record(record) for record in records if record.get("file_path")]


def normalize_library_record(record: dict) -> dict:
    title = record.get("title") or os.path.basename(record.get("file_path", "")) or record.get("paper_id", "")
    return {
        "paper_id": record.get("paper_id", ""),
        "name": title,
        "path": record.get("file_path", ""),
        "timestamp": _format_saved_at(record.get("saved_at", "")),
        "year": record.get("year") or 0,
        "page_count": record.get("page_count") or 0,
        "tags": record.get("tags") or [],
        "favorite": bool(record.get("favorite")),
        "status": record.get("status") or "已解析",
    }


def _format_saved_at(value: str) -> str:
    if not value:
        return ""
    try:
        return datetime.fromisoformat(value).strftime("%m-%d %H:%M")
    except ValueError:
        return str(value)[:16]


def add_record_to_current_session(record: dict) -> None:
    if not record.get("path"):
        return
    existing = next((p for p in st.session_state.uploaded_papers if p.get("paper_id") == record.get("paper_id")), None)
    payload = {
        "name": record.get("name", ""),
        "path": record.get("path", ""),
        "timestamp": record.get("timestamp") or datetime.now().strftime("%m-%d %H:%M"),
        "paper_id": record.get("paper_id", ""),
        "page_count": record.get("page_count", 0),
        "tags": record.get("tags", []),
        "favorite": record.get("favorite", False),
        "status": record.get("status", ""),
    }
    if existing:
        existing.update(payload)
    else:
        st.session_state.uploaded_papers.append(payload)


def update_library_tags(paper_id: str, tags: list[str]) -> None:
    st.session_state.tool_registry.execute(
        {
            "tool": "sqlite_paper_db",
            "kwargs": {"operation": "set_tags", "paper_id": paper_id, "tags": ",".join(tags)},
        }
    )


def set_library_favorite(paper_id: str, favorite: bool) -> None:
    st.session_state.tool_registry.execute(
        {"tool": "sqlite_paper_db", "kwargs": {"operation": "favorite" if favorite else "unfavorite", "paper_id": paper_id}}
    )


def parse_document(path: str | None) -> str:
    if not path or not os.path.exists(path):
        return ""
    cache_key = f"parsed::{path}::{os.path.getmtime(path)}"
    if cache_key in st.session_state:
        return st.session_state[cache_key]
    fmt = _detect_format(path)
    parser = PARSERS.get(fmt)
    if not parser:
        return f"[Unsupported format: {fmt}]"
    text = parser(path)
    st.session_state[cache_key] = text
    return text


def current_document_context(limit: int = 5000) -> str:
    path = st.session_state.get("current_doc")
    name = st.session_state.get("current_doc_name", "")
    text = parse_document(path)
    if not text or text.startswith("["):
        return ""
    return f"[Current document: {name}]\n{text[:limit]}"


def paper_record_from_path(path: str) -> dict:
    text = parse_document(path)
    sections = _extract_sections(text) if text and not text.startswith("[") else {}
    basename = os.path.basename(path)
    title = _clean(sections.get("title") or basename)
    method = _clean(sections.get("method") or _window(text, ["method", "approach", "framework", "model"]))
    results = _clean(sections.get("results") or _window(text, ["result", "experiment", "evaluation"]))
    return {
        "paper_id": basename,
        "title": title[:160] or basename,
        "method": method[:280],
        "dataset": _extract_mentions(text, ["dataset", "benchmark", "corpus"])[:180],
        "metrics": _extract_mentions(text, ["accuracy", "f1", "precision", "recall", "auc", "latency"])[:180],
        "results": results[:280],
        "key_finding": (results or method)[:280],
        "year": _guess_year(text),
    }


def uploaded_paper_records() -> list[dict]:
    return [
        paper_record_from_path(p["path"])
        for p in st.session_state.uploaded_papers
        if os.path.exists(p["path"])
    ]


def save_ui_session() -> None:
    if not st.session_state.messages:
        return
    session_id = st.session_state.get("current_session_id")
    if not session_id:
        session_id = datetime.now().strftime("%Y%m%d%H%M%S")
        st.session_state.current_session_id = session_id
    first_user = next((m["content"] for m in st.session_state.messages if m.get("role") == "user"), "新对话")
    st.session_state.sessions[session_id] = {
        "title": first_user[:24],
        "timestamp": datetime.now().strftime("%m-%d %H:%M"),
        "messages": list(st.session_state.messages),
    }


def run_agent_prompt(prompt: str, model_name: str | None = None) -> None:
    selected_model = normalize_model_choice(model_name)
    st.session_state.messages.append({"role": "user", "content": prompt, "model": selected_model})
    agent = get_agent(selected_model)
    if not agent:
        st.session_state.messages.append(
            {"role": "assistant", "content": "请先在“模型配置”页面添加并启用一个可用的大模型。", "model": selected_model}
        )
        save_ui_session()
        return
    try:
        result = agent.process_message(prompt, document_context=current_document_context())
        trace = result.get("trace") or result.get("推理过程") or []
        st.session_state.messages.append(
            {
                "role": "assistant",
                "content": result.get("response", "已完成，但没有生成可显示的回答。"),
                "trace": trace,
                "tools_called": result.get("tools_called", []),
                "model": selected_model,
            }
        )
    except Exception as exc:
        st.session_state.messages.append(
            {
                "role": "assistant",
                "content": f"运行时出现错误：`{type(exc).__name__}: {exc}`",
                "model": selected_model,
            }
        )
    save_ui_session()


def append_tool_result(title: str, content: str, model_name: str = NO_LLM, tool_name: str = "") -> None:
    st.session_state.messages.append(
        {
            "role": "assistant",
            "content": f"### {title}\n\n{content}",
            "model": model_name,
            "tools_called": [tool_name] if tool_name else [],
        }
    )
    save_ui_session()


def run_llm_structure_for_current_doc() -> None:
    path = st.session_state.get("current_doc")
    name = st.session_state.get("current_doc_name", "")
    text = parse_document(path)
    if not path or not text:
        st.session_state["last_structure_status"] = "请先选择论文。"
        return
    st.session_state[_doc_key(path, "llm_structure")] = llm_structure_document(text, name)
    st.session_state["last_structure_status"] = "结构化解析已生成。"


def run_full_translation_for_current_doc() -> None:
    path = st.session_state.get("current_doc")
    text = parse_document(path)
    if not path or not text:
        st.session_state["last_translation_status"] = "请先选择论文。"
        return
    st.session_state[_doc_key(path, "full_translation")] = translate_full_document(text)
    st.session_state["last_translation_status"] = "全文翻译已完成。"


def render_trace(trace: list[dict]) -> None:
    for step in trace:
        kind = step.get("type")
        if kind == "action":
            st.caption(f"步骤 {step.get('step')}: 调用工具 {step.get('tool')}")
            st.json(step.get("params", {}))
        elif kind == "observation":
            st.caption(f"步骤 {step.get('step')}: 工具返回")
            st.code(step.get("result", "")[:1200])
        elif kind == "thought":
            st.caption(f"步骤 {step.get('step')}: 思考")
            st.write(step.get("content", "")[:800])
        elif kind == "final_answer":
            st.caption(f"步骤 {step.get('step')}: 最终回答")


def render_message(message: dict) -> None:
    role = message.get("role", "assistant")
    model = message.get("model") or st.session_state.get("current_model", "")
    tools_called = message.get("tools_called") or []
    meta = [f'<span class="model-badge">{html.escape(model or NO_LLM)}</span>']
    if tools_called:
        meta.append(f"调用 {'、'.join(html.escape(str(t)) for t in tools_called)}")
    elif role == "user":
        meta.append("本条消息")
    else:
        meta.append("模型回答")
    st.markdown(
        f'<div class="message-card {html.escape(role)}"><div class="message-meta">{"".join(meta)}</div>',
        unsafe_allow_html=True,
    )
    st.markdown(message.get("content", ""))
    trace = message.get("trace")
    if trace:
        with st.expander("查看推理过程", expanded=False):
            render_trace(trace)
    st.markdown("</div>", unsafe_allow_html=True)


def render_library_card(paper: dict) -> None:
    paper_id = paper.get("paper_id") or paper.get("name", "")
    safe_key = re.sub(r"[^a-zA-Z0-9_]+", "_", paper_id or paper.get("name", "paper"))
    active = st.session_state.get("current_doc") == paper.get("path")
    tags = list(dict.fromkeys(paper.get("tags") or []))
    if paper.get("favorite") and "已收藏" not in tags:
        tags.insert(0, "已收藏")
    if not tags:
        tags = ["论文"]
    meta_left = f"{paper.get('page_count')} 页" if paper.get("page_count") else (str(paper.get("year")) if paper.get("year") else "本地论文")
    meta_right = "当前阅读" if active else (paper.get("status") or paper.get("timestamp") or "已入库")
    inner_class = "library-card-inner active" if active else "library-card-inner"
    with st.container(border=True):
        top_cols = st.columns([0.74, 0.26], vertical_alignment="top")
        with top_cols[0]:
            st.markdown(
                f'<div class="{inner_class}"><div class="library-card-title">{html.escape(paper.get("name", "未命名论文"))}</div></div>',
                unsafe_allow_html=True,
            )
        with top_cols[1]:
            menu = st.popover("...", key=f"paper_menu_{safe_key}", width="content", use_container_width=True)
        st.markdown(
            f"""
            <div class="library-meta-row"><span>{html.escape(str(meta_left))}</span><span>{html.escape(str(meta_right))}</span></div>
            <div class="library-tags">{''.join(render_library_tag(tag) for tag in tags[:4])}</div>
            """,
            unsafe_allow_html=True,
        )
        with menu:
            st.caption("论文操作")
            if st.button("打开论文", key=f"open_lib_{safe_key}", use_container_width=True):
                add_record_to_current_session(paper)
                st.session_state.current_doc = paper.get("path")
                st.session_state.current_doc_name = paper.get("name", "")
                st.rerun()
            favorite_label = "取消收藏" if paper.get("favorite") else "收藏"
            if st.button(favorite_label, key=f"fav_lib_{safe_key}", use_container_width=True):
                set_library_favorite(paper_id, not paper.get("favorite"))
                st.rerun()
            tags_text = st.text_input(
                "标签（逗号分隔）",
                value=", ".join(paper.get("tags") or []),
                key=f"tags_lib_{safe_key}",
            )
            if st.button("保存标签", key=f"save_tags_lib_{safe_key}", use_container_width=True):
                update_library_tags(paper_id, [tag.strip() for tag in re.split(r"[,，]", tags_text) if tag.strip()])
                st.rerun()
            if st.button("加入对比", key=f"compare_lib_{safe_key}", use_container_width=True):
                add_record_to_current_session(paper)
                compare_ids = st.session_state.setdefault("compare_paper_ids", [])
                if paper_id not in compare_ids:
                    compare_ids.append(paper_id)
                compare_tags = list(dict.fromkeys((paper.get("tags") or []) + ["待对比"]))
                update_library_tags(paper_id, compare_tags)
                st.rerun()
            if st.button("重新解析", key=f"reparse_lib_{safe_key}", use_container_width=True):
                for key in list(st.session_state.keys()):
                    if str(key).startswith(f"parsed::{paper.get('path')}"):
                        del st.session_state[key]
                st.session_state.current_doc = paper.get("path")
                st.session_state.current_doc_name = paper.get("name", "")
                st.rerun()
            if st.button("查看结构化信息", key=f"inspect_lib_{safe_key}", use_container_width=True):
                st.session_state.current_doc = paper.get("path")
                st.session_state.current_doc_name = paper.get("name", "")
                st.rerun()
            if st.button("从当前会话移除", key=f"remove_session_lib_{safe_key}", use_container_width=True):
                st.session_state.uploaded_papers = [
                    p for p in st.session_state.uploaded_papers if p.get("paper_id") != paper_id
                ]
                if active:
                    st.session_state.current_doc = None
                    st.session_state.current_doc_name = ""
                st.rerun()
            st.divider()
            confirm = st.checkbox("确认从论文库删除", key=f"confirm_delete_lib_{safe_key}")
            if st.button("从论文库删除", key=f"delete_lib_{safe_key}", use_container_width=True, disabled=not confirm):
                st.session_state.tool_registry.execute(
                    {"tool": "sqlite_paper_db", "kwargs": {"operation": "delete", "paper_id": paper_id}}
                )
                st.session_state.uploaded_papers = [
                    p for p in st.session_state.uploaded_papers if p.get("paper_id") != paper_id
                ]
                if active:
                    st.session_state.current_doc = None
                    st.session_state.current_doc_name = ""
                st.rerun()


def render_library_tag(tag: str) -> str:
    klass = "tag"
    if tag in {"已收藏", "已选择"}:
        klass += " teal"
    elif tag in {"待对比", "待处理"}:
        klass += " amber"
    return f'<span class="{klass}">{html.escape(tag)}</span>'


def render_sidebar() -> None:
    with st.sidebar:
        st.markdown(
            """
            <div class="sidebar-brand">
              <div class="sidebar-brand-mark">文</div>
              <div>
                <p class="sidebar-brand-title">学术文献分析助手</p>
                <p class="sidebar-brand-subtitle">Academic Lit Agent</p>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown(
            """
            <div class="asset-title-row">
              <h3>论文资产</h3>
              <div class="asset-add">＋</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        uploads = st.file_uploader(
            "上传 PDF / Word / Markdown",
            type=["pdf", "docx", "doc", "md", "txt"],
            accept_multiple_files=True,
            key="paper_uploads",
        )
        if uploads:
            for uploaded in uploads:
                add_uploaded_file(uploaded)
            st.success(f"已加入 {len(uploads)} 个文件")

        library_records = paper_library_records()
        if library_records:
            for paper in library_records:
                render_library_card(paper)
        else:
            st.caption("还没有上传论文。")

        st.divider()
        st.markdown("### 可用模型")
        models = active_model_names()
        if models:
            default_index = model_index(models, st.session_state.get("current_model"))
            st.session_state.current_model = st.selectbox("默认模型", models, index=default_index, key="global_model_select")
            if st.session_state.get("message_model") not in models:
                st.session_state.message_model = st.session_state.current_model
            if st.session_state.get("task_model") not in models:
                st.session_state.task_model = st.session_state.current_model
            for name in models:
                active = name == st.session_state.current_model
                st.markdown(
                    f"""
                    <div class="paper-chip {'active' if active else ''}">
                      <div class="paper-title"><span class="status-dot"></span> {html.escape(name)}</div>
                      <div class="paper-meta"><span>{'默认' if active else '在线'}</span><span>可用于任务</span></div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
        else:
            st.warning("还没有启用模型。")

        st.divider()
        st.markdown("### 工作区")
        st.caption("当前项目：数字孪生综述")
        st.caption(f"默认模型：{st.session_state.get('current_model') or '未配置'}")
        st.caption(f"模型策略：{st.session_state.get('model_policy', '允许单次切换')}")

        st.divider()
        st.markdown("### 快捷入口")
        if st.button("模型配置", use_container_width=True):
            st.switch_page("pages/model_config.py")
        if st.button("论文库统计", use_container_width=True):
            out = st.session_state.tool_registry.execute({"tool": "sqlite_paper_db", "kwargs": {"operation": "stats"}})
            append_tool_result("论文库统计", out, NO_LLM, "sqlite_paper_db")
            st.rerun()
        if st.button("去重论文库", use_container_width=True):
            out = st.session_state.tool_registry.execute({"tool": "sqlite_paper_db", "kwargs": {"operation": "dedupe"}})
            append_tool_result("论文库去重", out, NO_LLM, "sqlite_paper_db")
            st.rerun()
        if st.button("导出报告", use_container_width=True):
            append_tool_result("导出报告", "当前对话和分析结果已整理在分析流中。后续可以接入 Word/LaTeX 导出。", NO_LLM, "")
            st.rerun()

        st.divider()
        st.markdown("### 对话")
        if st.button("新建对话", use_container_width=True):
            st.session_state.messages = []
            st.session_state.current_session_id = None
            st.rerun()
        for sid, session in reversed(list(st.session_state.sessions.items())):
            if st.button(session["title"], key=f"session_{sid}", use_container_width=True):
                st.session_state.messages = session["messages"]
                st.session_state.current_session_id = sid
                st.rerun()


def render_topbar() -> None:
    left, default_col, policy_col, save_col, action_col = st.columns(
        [0.78, 0.66, 0.74, 0.34, 0.48],
        vertical_alignment="center",
    )
    with left:
        st.markdown(
            """
            <div class="topbar-title">
              <h1>文献研究工作台</h1>
              <p>默认模型管理全局任务；单条消息、工具调用和长任务可以临时指定模型。</p>
            </div>
            """,
            unsafe_allow_html=True,
        )
    models = active_model_names()
    with default_col:
        if models:
            idx = model_index(models, st.session_state.get("current_model"))
            st.session_state.current_model = st.selectbox(
                "默认模型",
                models,
                index=idx,
                key="top_default_model",
                label_visibility="collapsed",
                format_func=lambda name: f"默认： {name}",
            )
        else:
            st.selectbox(
                "默认模型",
                ["未配置"],
                disabled=True,
                label_visibility="collapsed",
                format_func=lambda name: f"默认： {name}",
            )
    with policy_col:
        policies = ["允许单次切换", "始终使用默认模型"]
        st.session_state.model_policy = st.selectbox(
            "任务策略",
            policies,
            index=model_index(policies, st.session_state.get("model_policy")),
            key="top_model_policy",
            label_visibility="collapsed",
            format_func=lambda policy: f"策略： {policy}",
        )
    with save_col:
        if st.button("保存会话", use_container_width=True):
            save_ui_session()
            st.success("已保存")
    with action_col:
        if st.button("生成综述草稿", type="primary", use_container_width=True):
            st.session_state.pending_prompt = {
                "text": "请基于当前论文、分析流和已上传论文，生成一份文献综述草稿提纲，包含研究背景、方法对比、研究空白和可写作小节。",
                "model": normalize_model_choice(st.session_state.get("task_model")),
            }
            st.rerun()


def render_metrics() -> None:
    metric_cols = st.columns(4)
    metric_cols[0].metric("当前论文", len(st.session_state.uploaded_papers), "已上传")
    metric_cols[1].metric("已调用工具", len(st.session_state.tool_registry.list_tools()), "全部可用")
    metric_cols[2].metric("当前消息", len(st.session_state.messages), "分析流")
    metric_cols[3].metric("本条模型", st.session_state.get("message_model") or "未配置", "可临时切换")


def render_tool_cards() -> None:
    st.markdown('<div class="panel-head"><h3>分析流</h3><span>每次工具调用都会记录使用的模型</span></div>', unsafe_allow_html=True)
    st.markdown('<div style="padding: 12px 14px; background: #fbfcfe; border-bottom: 1px solid #ebeff5;">', unsafe_allow_html=True)
    cards = st.columns(4)
    labels = [
        ("搜索论文", "关键词、年份、领域和分页", "search_model", True, True),
        ("多论文对比", "方法、数据集、指标、结果", "compare_model", True, False),
        ("研究空白", "共识、矛盾、空白、趋势", "gap_model", True, False),
        ("引用图", "BFS / DFS 与共现关系", "citation_model", True, True),
    ]
    for col, (title, desc, key, include_auto, include_no_llm) in zip(cards, labels):
        with col:
            st.markdown(f'<div class="tool-card"><strong>{title}</strong><p>{desc}</p>', unsafe_allow_html=True)
            options = model_options(include_auto=include_auto, include_no_llm=include_no_llm)
            default = AUTO_MODEL if include_auto else st.session_state.get("task_model")
            if title == "引用图":
                default = NO_LLM
            st.selectbox("模型", options, index=model_index(options, st.session_state.get(key, default)), key=key, label_visibility="collapsed")
            st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)


def render_tool_console() -> None:
    tools = st.session_state.tool_registry
    with st.expander("展开工具参数", expanded=False):
        tab_search, tab_analysis, tab_citation, tab_code = st.tabs(["搜索", "分析", "引用", "代码"])
        with tab_search:
            q = st.text_input("关键词", placeholder="retrieval augmented generation")
            cols = st.columns([1, 1, 1, 2])
            limit = cols[0].number_input("数量", 1, 50, 5)
            offset = cols[1].number_input("分页", 0, 500, 0, step=5)
            year = cols[2].text_input("年份", placeholder="2020-2026")
            field = cols[3].text_input("领域", placeholder="Computer Science")
            if st.button("搜索论文", use_container_width=True):
                if not q.strip():
                    st.warning("请输入关键词。")
                else:
                    out = tools.execute(
                        {
                            "tool": "search_papers",
                            "kwargs": {
                                "query": q,
                                "limit": limit,
                                "offset": offset,
                                "year": year,
                                "fields_of_study": field,
                            },
                        }
                    )
                    append_tool_result("论文搜索结果", out, st.session_state.get("search_model", AUTO_MODEL), "search_papers")
                    st.rerun()
        with tab_analysis:
            cols = st.columns(3)
            if cols[0].button("总结当前论文", use_container_width=True, disabled=not st.session_state.get("current_doc")):
                st.session_state.pending_prompt = {
                    "text": "请阅读当前论文，概括标题、研究问题、方法、实验结果和主要贡献。",
                    "model": normalize_model_choice(st.session_state.get("task_model")),
                }
                st.rerun()
            if cols[1].button("对比已上传论文", use_container_width=True, disabled=len(st.session_state.uploaded_papers) < 2):
                paths = [p["path"] for p in st.session_state.uploaded_papers]
                out = tools.execute(
                    {
                        "tool": "literature_comparator",
                        "kwargs": {"paper_paths_json": json.dumps(paths, ensure_ascii=False), "topic": "已上传论文对比"},
                    }
                )
                append_tool_result("多论文对比", out, st.session_state.get("compare_model", AUTO_MODEL), "literature_comparator")
                st.rerun()
            if cols[2].button("研究空白分析", use_container_width=True, disabled=len(st.session_state.uploaded_papers) < 2):
                records = uploaded_paper_records()
                out = tools.execute({"tool": "research_gap_analyzer", "kwargs": {"papers_json": json.dumps(records, ensure_ascii=False)}})
                append_tool_result("研究空白分析", out, st.session_state.get("gap_model", AUTO_MODEL), "research_gap_analyzer")
                st.rerun()
        with tab_citation:
            cid = st.text_input("论文 ID", placeholder="ARXIV:1706.03762")
            other = st.text_input("共现论文 ID", placeholder="可选")
            cols = st.columns(3)
            direction = cols[0].selectbox("方向", ["citations", "references"])
            depth = cols[1].number_input("深度", 1, 2, 1)
            citation_limit = cols[2].number_input("数量", 1, 50, 10, key="citation_limit")
            if st.button("运行引用分析", use_container_width=True):
                if not cid.strip():
                    st.warning("请输入论文 ID。")
                else:
                    kwargs = {"paper_id": cid, "direction": direction, "depth": depth, "limit": citation_limit}
                    if other:
                        kwargs.update({"operation": "co_occurrence", "other_paper_id": other})
                    out = tools.execute({"tool": "citation_traverser", "kwargs": kwargs})
                    append_tool_result("引用分析", out, st.session_state.get("citation_model", NO_LLM), "citation_traverser")
                    st.rerun()
        with tab_code:
            code = st.text_area("Python 代码", value="print('hello literature')", height=160)
            if st.button("执行代码", use_container_width=True):
                out = tools.execute({"tool": "code_executor", "kwargs": {"code": code, "timeout": 8}})
                append_tool_result("代码执行结果", out, NO_LLM, "code_executor")
                st.rerun()


def render_chat_feed() -> None:
    feed = st.container(height=180, border=False)
    with feed:
        if not st.session_state.messages:
            st.info("可以直接提问，也可以从工具参数里启动搜索、对比、研究空白或引用分析。")
        for message in st.session_state.messages:
            render_message(message)

    models = active_model_names()
    with st.form("message_form", clear_on_submit=True, border=False):
        cols = st.columns([0.07, 0.55, 0.25, 0.13], vertical_alignment="center")
        with cols[0]:
            st.markdown('<div class="composer-plus">＋</div>', unsafe_allow_html=True)
        with cols[1]:
            prompt = st.text_area(
                "消息",
                placeholder="继续提问，或输入 / 调用工具...",
                height=52,
                label_visibility="collapsed",
            )
        with cols[2]:
            if models:
                current = st.session_state.get("message_model") or st.session_state.get("current_model")
                selected = st.selectbox(
                    "本条模型",
                    models,
                    index=model_index(models, current),
                    key="message_model_select",
                    label_visibility="collapsed",
                )
            else:
                selected = ""
                st.selectbox("本条模型", ["未配置"], disabled=True, label_visibility="collapsed")
        with cols[3]:
            submitted = st.form_submit_button("发送", type="primary", use_container_width=True)
        if submitted and prompt.strip():
            st.session_state.message_model = selected
            with st.spinner("正在思考并调用工具..."):
                run_agent_prompt(prompt.strip(), selected)
            st.rerun()


def render_document_preview(path: str | None, name: str) -> None:
    st.markdown(
        f'<div class="panel-head"><h3>阅读与结构化信息</h3><span>摘要模型：{html.escape(st.session_state.get("current_model") or "未配置")}</span></div>',
        unsafe_allow_html=True,
    )
    if not path or not os.path.exists(path):
        st.info("上传或选择论文后，这里会显示原文预览和解析内容。")
        return

    fmt = _detect_format(path)
    text = parse_document(path)
    toolbar = st.columns([1, 1, 1, 2])
    with toolbar[0]:
        if st.button("重新解析", use_container_width=True):
            for key in list(st.session_state.keys()):
                if str(key).startswith(f"parsed::{path}"):
                    del st.session_state[key]
            st.rerun()
    with toolbar[1]:
        st.button("LLM解析", use_container_width=True, key="llm_structure_top", on_click=run_llm_structure_for_current_doc)
    with toolbar[2]:
        st.button("翻译全文", use_container_width=True, key="translate_full_top", on_click=run_full_translation_for_current_doc)
    with toolbar[3]:
        st.caption(f"{name} · {fmt.upper()}")

    preview_tab, structure_tab, translation_tab, notes_tab = st.tabs(["原文", "结构化解析", "翻译", "笔记"])
    with preview_tab:
        if fmt == "pdf" and _has_lib("fitz"):
            render_pdf_page(path)
        elif fmt == "docx" and _has_lib("docx"):
            render_docx(path)
        elif fmt in {"md", "txt"}:
            st.text_area("文本预览", text[:10000], height=520, label_visibility="collapsed")
        else:
            st.warning("当前环境缺少该格式的预览依赖，已显示解析文本。")
            st.text_area("解析文本", text[:10000], height=520, label_visibility="collapsed")
    with structure_tab:
        sections = _extract_sections(text) if text and not text.startswith("[") else {}
        if not st.session_state.get(_doc_key(path, "llm_structure"), ""):
            st.button("调用 LLM 生成结构化解析", use_container_width=True, key="llm_structure_tab", on_click=run_llm_structure_for_current_doc)
        if st.session_state.get("last_structure_status"):
            st.caption(st.session_state["last_structure_status"])
        render_document_inspector(sections, text, path)
    with translation_tab:
        translation = st.session_state.get(_doc_key(path, "full_translation"), "")
        if translation:
            with st.container(height=640, border=True):
                st.markdown(translation)
        else:
            st.caption("点击“翻译全文”后显示完整译文。")
    with notes_tab:
        st.text_area("研究笔记", placeholder="记录综述线索、实验设置、可引用结论...", height=420)


def render_document_inspector(sections: dict, text: str, path: str) -> None:
    if text.startswith("["):
        st.warning(text)
        return
    llm_structure = st.session_state.get(_doc_key(path, "llm_structure"), "")
    if llm_structure:
        st.markdown("#### LLM 结构化解析")
        st.markdown(llm_structure)
        st.divider()
    else:
        st.info("点击上方“LLM解析”可调用当前模型生成结构化解析。下方为规则解析结果。")
    left, right = st.columns([1, 0.7])
    with left:
        for key, label in [
            ("title", "标题"),
            ("abstract", "摘要"),
            ("method", "方法"),
            ("results", "实验与结果"),
            ("conclusion", "结论"),
            ("references", "参考文献"),
        ]:
            value = sections.get(key, "").strip()
            if value:
                with st.expander(label, expanded=key in {"title", "abstract"}):
                    st.write(value[:3000])
    with right:
        st.markdown(
            """
            <div class="inspector-card">
              <h4>识别出的综述线索</h4>
              <div class="hint-box teal">共识：数字孪生可改善系统感知与反馈。</div>
              <div class="hint-box">空白：真实部署、跨域数据与成本评估不足。</div>
              <div class="hint-box">矛盾：效率提升与模型个性化之间存在取舍。</div>
            </div>
            """,
            unsafe_allow_html=True,
        )


def render_pdf_page(path: str) -> None:
    import fitz

    cache_key = _doc_key(path, "pdf_pages")
    if cache_key not in st.session_state:
        doc = fitz.open(path)
        try:
            pages = []
            for page_index in range(len(doc)):
                pix = doc[page_index].get_pixmap(matrix=fitz.Matrix(1.15, 1.15))
                pages.append(pix.tobytes("png"))
            st.session_state[cache_key] = pages
        finally:
            doc.close()
    pages = st.session_state[cache_key]
    with st.container(height=640, border=False):
        for index, image in enumerate(pages, 1):
            st.caption(f"第 {index} 页 / 共 {len(pages)} 页")
            st.image(image, use_container_width=True)


def render_docx(path: str) -> None:
    import docx

    doc = docx.Document(path)
    shown = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            st.markdown(text)
            shown += 1
        if shown >= 80:
            st.caption("已显示前 80 段。")
            break


def translate_current_summary(text: str) -> None:
    client = get_llm_client(st.session_state.get("current_model"))
    if not client:
        st.session_state.doc_translation = "请先配置大模型。"
        return
    sections = _extract_sections(text)
    source = "\n\n".join(
        part
        for part in [sections.get("title", ""), sections.get("abstract", ""), sections.get("conclusion", "")]
        if part
    )
    if not source:
        source = text[:2500]
    try:
        st.session_state.doc_translation = client.chat(
            [
                {
                    "role": "user",
                    "content": "将下面的学术文本翻译为中文，保留术语、引用和公式，只输出译文。\n\n" + source[:3000],
                }
            ],
            temperature=0.1,
        )
    except Exception as exc:
        st.session_state.doc_translation = f"翻译失败：{exc}"


def llm_structure_document(text: str, name: str) -> str:
    client = get_llm_client(st.session_state.get("current_model"))
    if not client:
        return "请先配置大模型。"
    sections = _extract_sections(text)
    source = "\n\n".join(
        [
            f"文件名：{name}",
            "标题：\n" + sections.get("title", ""),
            "摘要：\n" + sections.get("abstract", ""),
            "方法：\n" + sections.get("method", ""),
            "实验与结果：\n" + sections.get("results", ""),
            "结论：\n" + sections.get("conclusion", ""),
            "原文片段：\n" + text[:4000],
        ]
    )
    prompt = (
        "请对下面论文内容做结构化解析，用中文输出 Markdown。必须包含：\n"
        "1. 标题\n2. 研究问题\n3. 核心方法\n4. 数据集/场景\n5. 评价指标\n"
        "6. 关键结果\n7. 主要贡献\n8. 局限性\n9. 可用于文献综述的写作线索。\n"
        "如果信息缺失，请写“原文未明确”。不要编造。\n\n"
        + source[:8000]
    )
    try:
        result = client.chat([{"role": "user", "content": prompt}], temperature=0.1, max_tokens=1800)
        if result and result.strip():
            return result
        return _fallback_structure(sections)
    except Exception as exc:
        return f"LLM 结构化解析失败：{exc}\n\n{_fallback_structure(sections)}"


def _fallback_structure(sections: dict) -> str:
    return "\n\n".join(
        [
            "> LLM 未返回有效文本，以下为规则解析兜底结果。",
            "### 标题\n" + (sections.get("title", "原文未明确")[:800] or "原文未明确"),
            "### 研究问题\n" + _clean(sections.get("abstract", "原文未明确"))[:700],
            "### 核心方法\n" + _clean(sections.get("method", "原文未明确"))[:700],
            "### 关键结果\n" + _clean(sections.get("results", "原文未明确"))[:700],
            "### 局限性与综述线索\n原文未明确，需要结合更多论文继续分析。",
        ]
    )


def translate_full_document(text: str) -> str:
    client = get_llm_client(st.session_state.get("current_model"))
    if not client:
        return "请先配置大模型。"
    clean_text = _clean_for_translation(text)
    chunks = _split_text(clean_text, 2800)
    if not chunks:
        return "没有可翻译的文本。"
    progress = st.progress(0, text="准备翻译全文")
    translated = []
    for index, chunk in enumerate(chunks, 1):
        progress.progress((index - 1) / len(chunks), text=f"正在翻译第 {index}/{len(chunks)} 段")
        prompt = (
            "将下面的学术论文正文完整翻译为中文。要求：\n"
            "- 保留术语、公式、引用编号和小节编号。\n"
            "- 不要总结，不要省略，不要添加解释。\n"
            "- 只输出译文。\n\n"
            f"第 {index}/{len(chunks)} 段：\n{chunk}"
        )
        try:
            translated.append(client.chat([{"role": "user", "content": prompt}], temperature=0.1, max_tokens=3500))
        except Exception as exc:
            translated.append(f"\n\n[第 {index} 段翻译失败：{exc}]\n\n{chunk}")
    progress.progress(1.0, text="全文翻译完成")
    return "\n\n".join(translated)


def _split_text(text: str, chunk_size: int) -> list[str]:
    chunks = []
    current = []
    current_len = 0
    for paragraph in re.split(r"\n{2,}", text):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        parts = [paragraph]
        if len(paragraph) > chunk_size:
            parts = [paragraph[i:i + chunk_size] for i in range(0, len(paragraph), chunk_size)]
        for part in parts:
            if current and current_len + len(part) > chunk_size:
                chunks.append("\n\n".join(current))
                current = [part]
                current_len = len(part)
            else:
                current.append(part)
                current_len += len(part)
    if current:
        chunks.append("\n\n".join(current))
    return chunks


def _clean_for_translation(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    cleaned = []
    for line in lines:
        if not line:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def _doc_key(path: str, suffix: str) -> str:
    try:
        stamp = os.path.getmtime(path)
    except OSError:
        stamp = 0
    return f"doc::{suffix}::{path}::{stamp}"


def _has_lib(name: str) -> bool:
    try:
        __import__(name)
        return True
    except ImportError:
        return False


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _window(text: str, needles: list[str], size: int = 900) -> str:
    lower = (text or "").lower()
    positions = [lower.find(n) for n in needles if lower.find(n) >= 0]
    if not positions:
        return (text or "")[:size]
    start = max(0, min(positions) - 120)
    return text[start:start + size]


def _extract_mentions(text: str, hints: list[str]) -> str:
    lower = (text or "").lower()
    found = [hint.upper() if hint in {"f1", "auc"} else hint for hint in hints if hint in lower]
    return ", ".join(dict.fromkeys(found))


def _guess_year(text: str) -> int:
    years = [int(y) for y in re.findall(r"\b(20[0-3][0-9]|19[8-9][0-9])\b", text or "")]
    return max(years) if years else 0


ensure_tool_registry()
page_css()
render_sidebar()
render_topbar()
render_metrics()

pending_prompt = st.session_state.pop("pending_prompt", None)
if pending_prompt:
    with st.spinner("正在分析..."):
        if isinstance(pending_prompt, dict):
            run_agent_prompt(pending_prompt["text"], pending_prompt.get("model"))
        else:
            run_agent_prompt(str(pending_prompt), st.session_state.get("task_model"))
    st.rerun()

chat_col, doc_col = st.columns([1.08, 0.92], gap="medium")
with chat_col:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    render_tool_cards()
    render_tool_console()
    render_chat_feed()
    st.markdown("</div>", unsafe_allow_html=True)

with doc_col:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    render_document_preview(st.session_state.get("current_doc"), st.session_state.get("current_doc_name", ""))
    st.markdown("</div>", unsafe_allow_html=True)
